#!/usr/bin/env python3
"""
DOCX <-> Markdown bridge for fixed-style technical reports.

Supported structure:
- Heading 1, Heading 2, Normal paragraphs.
- Inline/display equations from OMML to LaTeX-like math text.

Roundtrip strategy for equations:
- DOCX -> Markdown: emit LaTeX plus optional hidden OMML payload markers.
- Markdown -> DOCX: restore equation objects from payload markers when present.
  If markers are absent, keep equations as literal $...$ text.
"""

from __future__ import annotations

import argparse
import base64
import json
import re
import sys
import zlib
from collections import Counter
from dataclasses import dataclass
from hashlib import sha256
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Tuple, Union, cast
from xml.sax.saxutils import escape as xml_escape
from zipfile import ZipFile

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, RGBColor
from docx.styles.style import ParagraphStyle
import lxml.etree as etree


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
NS = {"w": W_NS, "m": M_NS}

INLINE_MARKER_RE = re.compile(r"<!--OMML_INLINE_Z:([A-Za-z0-9+/=]+)-->")
BLOCK_MARKER_RE = re.compile(r"^\s*<!--OMML_BLOCK_Z:([A-Za-z0-9+/=]+)-->\s*$")
PARAGRAPH_STYLE_MARKER_RE = re.compile(r"^\s*<!--DOCX_PSTYLE:([A-Za-z0-9+/=]+)-->\s*$")
FIRST_LINE_INDENT_MARKER_RE = re.compile(r"^\s*<!--DOCX_FIRST_LINE_INDENT_CM:([0-9]+(?:\.[0-9]+)?)-->\s*$")
PARAGRAPH_ALIGN_MARKER_RE = re.compile(r"^\s*<!--DOCX_PALIGN:([A-Za-z]+)-->\s*$")
TABLE_META_MARKER_RE = re.compile(r"^\s*<!--DOCX_TABLE_META:([A-Za-z0-9+/=]+)-->\s*$")
EMPTY_PARAGRAPH_MARKER = "<!--EMPTY_P-->"
DEFAULT_FIRST_LINE_INDENT_CM = 0.74
TWIPS_PER_CM = 1440 / 2.54
BR_TAG_RE = re.compile(r"<br\s*/?>", re.IGNORECASE)
DOUBLE_BR_TAG_RE = re.compile(r"(?:<br\s*/?>){2}", re.IGNORECASE)
COLOR_SPAN_RE = re.compile(r"<span\b(?P<attrs>[^>]*)>(?P<content>.*?)</span>", re.IGNORECASE | re.DOTALL)
COLOR_STYLE_RE = re.compile(r"(?:^|;)\s*color\s*:\s*(#[0-9A-Fa-f]{6}|[0-9A-Fa-f]{6})\s*(?:;|$)", re.IGNORECASE)
BACKGROUND_STYLE_RE = re.compile(r"(?:^|;)\s*background-color\s*:\s*(#[0-9A-Fa-f]{6}|[0-9A-Fa-f]{6})\s*(?:;|$)", re.IGNORECASE)
UNDERLINE_TAG_RE = re.compile(r"<u\b(?P<attrs>[^>]*)>(?P<content>.*?)</u>", re.IGNORECASE | re.DOTALL)
STRIKE_RE = re.compile(r"~~(.+?)~~", re.DOTALL)
HTML_ATTR_RE = re.compile(r"(?P<name>[A-Za-z_:][A-Za-z0-9_:\-]*)\s*=\s*(['\"])(?P<value>.*?)\2", re.DOTALL)

TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")
RUN_FONT_ATTRS = ("ascii", "hAnsi", "eastAsia", "cs")
WORD_HIGHLIGHT_TO_HEX = {
    "yellow": "FFFF00",
    "green": "00FF00",
    "cyan": "00FFFF",
    "magenta": "FF00FF",
    "blue": "0000FF",
    "red": "FF0000",
    "darkBlue": "000080",
    "darkCyan": "008080",
    "darkGreen": "008000",
    "darkMagenta": "800080",
    "darkRed": "800000",
    "darkYellow": "808000",
    "darkGray": "808080",
    "lightGray": "C0C0C0",
    "black": "000000",
}


@dataclass
class MathToken:
    latex: str
    omml_payload: Optional[str]
    block: bool = False


@dataclass
class TextToken:
    text: str
    bold: bool = False
    italic: bool = False
    color: Optional[str] = None
    underline: Optional[str] = None
    strike: bool = False
    background_color: Optional[str] = None
    highlight: Optional[str] = None
    character_style: Optional[str] = None


@dataclass
class InlineStyleToken:
    color: Optional[str] = None
    underline: Optional[str] = None
    strike: bool = False
    background_color: Optional[str] = None
    highlight: Optional[str] = None
    character_style: Optional[str] = None


Segment = Union[TextToken, MathToken]
MarkdownInlineSegment = Union[str, MathToken]


def xpath_elements(node: etree._Element, expression: str) -> List[etree._Element]:
    return cast(List[etree._Element], node.xpath(expression, namespaces=NS))


def xpath_strings(node: etree._Element, expression: str) -> List[str]:
    values = cast(List[Any], node.xpath(expression, namespaces=NS))
    return [value for value in values if isinstance(value, str)]


def wtag(name: str) -> str:
    return f"{{{W_NS}}}{name}"


def mtag(name: str) -> str:
    return f"{{{M_NS}}}{name}"


def local_name(node: etree._Element) -> str:
    return etree.QName(node).localname if isinstance(node.tag, str) else ""


def read_xml_from_docx(docx_path: Path) -> tuple[etree._Element, etree._Element]:
    with ZipFile(docx_path) as archive:
        document_xml = archive.read("word/document.xml")
        styles_xml = archive.read("word/styles.xml")
    return etree.fromstring(document_xml), etree.fromstring(styles_xml)


def style_id_to_name(styles_root: etree._Element) -> Dict[str, str]:
    mapping: Dict[str, str] = {}
    for style in xpath_elements(styles_root, ".//w:style[@w:type='paragraph']"):
        style_id = style.get(wtag("styleId"), "")
        name = xpath_strings(style, "./w:name/@w:val")
        mapping[style_id] = name[0] if name else style_id
    return mapping


def font_spec_from_node(node: etree._Element) -> Optional[Dict[str, str]]:
    spec = {
        attr: value
        for attr in RUN_FONT_ATTRS
        if (value := node.get(wtag(attr)))
    }
    return spec or None


def infer_default_run_font(docx_path: Path) -> Optional[Dict[str, str]]:
    document_root, styles_root = read_xml_from_docx(docx_path)

    font_counter: Counter[Tuple[Optional[str], Optional[str], Optional[str], Optional[str]]] = Counter()
    for node in xpath_elements(document_root, "/w:document/w:body//w:rPr/w:rFonts"):
        font_values = (
            node.get(wtag("ascii")),
            node.get(wtag("hAnsi")),
            node.get(wtag("eastAsia")),
            node.get(wtag("cs")),
        )
        if any(font_values):
            font_counter[font_values] += 1

    if font_counter:
        values = font_counter.most_common(1)[0][0]
        return {attr: value for attr, value in zip(RUN_FONT_ATTRS, values) if value}

    fallback_queries = (
        "/w:styles/w:style[@w:type='paragraph' and @w:styleId='Normal']/w:rPr/w:rFonts",
        "/w:styles/w:style[@w:type='paragraph' and ./w:name/@w:val='Normal']/w:rPr/w:rFonts",
        "/w:styles/w:docDefaults/w:rPrDefault/w:rPr/w:rFonts",
    )
    for query in fallback_queries:
        nodes = xpath_elements(styles_root, query)
        for node in nodes:
            spec = font_spec_from_node(node)
            if spec:
                return spec

    return None


def apply_font_spec_to_run(run, font_spec: Optional[Dict[str, str]]) -> None:
    if not font_spec:
        return

    font_name = font_spec.get("ascii") or font_spec.get("hAnsi")
    if font_name:
        run.font.name = font_name

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = etree.Element(qn("w:rFonts"))
        r_pr.append(r_fonts)

    for attr in RUN_FONT_ATTRS:
        value = font_spec.get(attr)
        if value:
            r_fonts.set(qn(f"w:{attr}"), value)


def normalize_color_value(color: Optional[str]) -> Optional[str]:
    if not color:
        return None

    normalized = color.strip()
    if normalized.startswith("#"):
        normalized = normalized[1:]
    if not re.fullmatch(r"[0-9A-Fa-f]{6}", normalized):
        return None
    return normalized.upper()


def normalize_underline_value(underline: Optional[str]) -> Optional[str]:
    if not underline:
        return None
    normalized = underline.strip()
    if not normalized or normalized.lower() == "none":
        return None
    return normalized


def normalize_highlight_value(highlight: Optional[str]) -> Optional[str]:
    if not highlight:
        return None
    normalized = highlight.strip()
    if not normalized:
        return None
    for key in WORD_HIGHLIGHT_TO_HEX:
        if key.lower() == normalized.lower():
            return key
    return normalized


def extract_html_attr(attrs: str, attr_name: str) -> Optional[str]:
    for match in HTML_ATTR_RE.finditer(attrs):
        if match.group("name").lower() == attr_name.lower():
            return match.group("value")
    return None


def apply_text_color_to_run(run, color: Optional[str]) -> None:
    normalized = normalize_color_value(color)
    if not normalized:
        return
    run.font.color.rgb = RGBColor.from_string(normalized)


def apply_run_underline_to_run(run, underline: Optional[str]) -> None:
    normalized = normalize_underline_value(underline)
    if not normalized:
        return

    r_pr = run._element.get_or_add_rPr()
    underline_node = r_pr.find(qn("w:u"))
    if underline_node is None:
        underline_node = etree.Element(qn("w:u"))
        r_pr.append(underline_node)
    underline_node.set(qn("w:val"), normalized)


def apply_run_strike_to_run(run, strike: bool) -> None:
    if not strike:
        return
    run.font.strike = True


def apply_run_background_to_run(
    run,
    background_color: Optional[str] = None,
    highlight: Optional[str] = None,
) -> None:
    normalized_background = normalize_color_value(background_color)
    normalized_highlight = normalize_highlight_value(highlight)
    r_pr = run._element.get_or_add_rPr()

    if normalized_highlight:
        highlight_node = r_pr.find(qn("w:highlight"))
        if highlight_node is None:
            highlight_node = etree.Element(qn("w:highlight"))
            r_pr.append(highlight_node)
        highlight_node.set(qn("w:val"), normalized_highlight)

    mapped_highlight_bg = WORD_HIGHLIGHT_TO_HEX.get(normalized_highlight or "")
    if normalized_background and normalized_background != mapped_highlight_bg:
        shd = r_pr.find(qn("w:shd"))
        if shd is None:
            shd = etree.Element(qn("w:shd"))
            r_pr.append(shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), normalized_background)


def find_character_style(doc: DocxDocument, style_id: Optional[str]) -> Optional[ParagraphStyle]:
    if not style_id:
        return None

    for style in doc.styles:
        if style.type != WD_STYLE_TYPE.CHARACTER:
            continue
        if style.style_id == style_id or style.name == style_id:
            return cast(ParagraphStyle, style)
    return None


def apply_character_style_to_run(run, doc: Optional[DocxDocument], character_style: Optional[str]) -> None:
    if doc is None or not character_style:
        return
    style = find_character_style(doc, character_style)
    if style is not None:
        run.style = style


def run_color_value(run: etree._Element) -> Optional[str]:
    colors = xpath_strings(run, "./w:rPr/w:color/@w:val")
    if not colors:
        return None

    color = normalize_color_value(colors[0])
    return color


def run_underline_value(run: etree._Element) -> Optional[str]:
    underline_nodes = xpath_elements(run, "./w:rPr/w:u")
    if not underline_nodes:
        return None
    return normalize_underline_value(underline_nodes[0].get(wtag("val")) or "single")


def run_strike_value(run: etree._Element) -> bool:
    strike_nodes = xpath_elements(run, "./w:rPr/w:strike")
    if not strike_nodes:
        return False
    val = strike_nodes[0].get(wtag("val"))
    if val is None:
        return True
    return val.lower() not in {"0", "false", "off"}


def run_highlight_value(run: etree._Element) -> Optional[str]:
    highlights = xpath_strings(run, "./w:rPr/w:highlight/@w:val")
    if not highlights:
        return None
    return normalize_highlight_value(highlights[0])


def run_background_color_value(run: etree._Element) -> Optional[str]:
    fills = xpath_strings(run, "./w:rPr/w:shd/@w:fill")
    if fills:
        fill = fills[0]
        if fill.lower() != "auto":
            normalized_fill = normalize_color_value(fill)
            if normalized_fill:
                return normalized_fill

    highlight = run_highlight_value(run)
    if highlight and highlight in WORD_HIGHLIGHT_TO_HEX:
        return WORD_HIGHLIGHT_TO_HEX[highlight]

    return None


def run_character_style_value(run: etree._Element) -> Optional[str]:
    styles = xpath_strings(run, "./w:rPr/w:rStyle/@w:val")
    return styles[0] if styles else None


def inline_style_from_span_attrs(attrs: str) -> Optional[InlineStyleToken]:
    style_text = extract_html_attr(attrs, "style") or ""
    color = None
    background_color = None

    color_match = COLOR_STYLE_RE.search(style_text)
    if color_match:
        color = normalize_color_value(color_match.group(1))

    background_match = BACKGROUND_STYLE_RE.search(style_text)
    if background_match:
        background_color = normalize_color_value(background_match.group(1))

    highlight = normalize_highlight_value(extract_html_attr(attrs, "data-docx-highlight"))
    if highlight and not background_color:
        background_color = WORD_HIGHLIGHT_TO_HEX.get(highlight)

    character_style = extract_html_attr(attrs, "data-docx-rstyle")
    if character_style:
        character_style = character_style.strip() or None

    if not any((color, background_color, highlight, character_style)):
        return None

    return InlineStyleToken(
        color=color,
        background_color=background_color,
        highlight=highlight,
        character_style=character_style,
    )


def split_style_spans(text: str) -> List[Tuple[str, InlineStyleToken]]:
    spans: List[Tuple[str, InlineStyleToken]] = []
    pos = 0

    for match in COLOR_SPAN_RE.finditer(text):
        inline_style = inline_style_from_span_attrs(match.group("attrs"))
        if inline_style is None:
            continue

        if match.start() > pos:
            spans.append((text[pos : match.start()], InlineStyleToken()))
        spans.append((match.group("content"), inline_style))
        pos = match.end()

    if pos < len(text):
        spans.append((text[pos:], InlineStyleToken()))
    if not spans:
        spans.append((text, InlineStyleToken()))
    return spans


def underline_value_from_attrs(attrs: str) -> Optional[str]:
    return normalize_underline_value(extract_html_attr(attrs, "data-docx-underline") or "single")


def split_underline_spans(text: str) -> List[Tuple[str, Optional[str]]]:
    spans: List[Tuple[str, Optional[str]]] = []
    pos = 0

    for match in UNDERLINE_TAG_RE.finditer(text):
        underline = underline_value_from_attrs(match.group("attrs"))
        if underline is None:
            continue

        if match.start() > pos:
            spans.append((text[pos : match.start()], None))
        spans.append((match.group("content"), underline))
        pos = match.end()

    if pos < len(text):
        spans.append((text[pos:], None))
    if not spans:
        spans.append((text, None))
    return spans


def split_strike_spans(text: str) -> List[Tuple[str, bool]]:
    spans: List[Tuple[str, bool]] = []
    pos = 0

    for match in STRIKE_RE.finditer(text):
        if match.start() > pos:
            spans.append((text[pos : match.start()], False))
        spans.append((match.group(1), True))
        pos = match.end()

    if pos < len(text):
        spans.append((text[pos:], False))
    if not spans:
        spans.append((text, False))
    return spans


def format_marker_number(value: float) -> str:
    return f"{value:.2f}".rstrip("0").rstrip(".")


def encode_paragraph_style_marker(style_id: str, style_name: str) -> str:
    payload = base64.b64encode(
        json.dumps(
            {"style_id": style_id, "style_name": style_name},
            ensure_ascii=False,
            separators=(",", ":"),
        ).encode("utf-8")
    ).decode("ascii")
    return f"<!--DOCX_PSTYLE:{payload}-->"


def encode_first_line_indent_marker(indent_cm: float) -> str:
    return f"<!--DOCX_FIRST_LINE_INDENT_CM:{format_marker_number(indent_cm)}-->"


def encode_paragraph_alignment_marker(alignment: str) -> str:
    return f"<!--DOCX_PALIGN:{alignment}-->"


def encode_table_meta_marker(table_meta: Dict[str, Any]) -> str:
    raw = json.dumps(table_meta, ensure_ascii=False, separators=(",", ":"), sort_keys=True).encode("utf-8")
    payload = base64.b64encode(zlib.compress(raw, level=9)).decode("ascii")
    return f"<!--DOCX_TABLE_META:{payload}-->"


def decode_paragraph_style_marker(line: str) -> Optional[Dict[str, str]]:
    match = PARAGRAPH_STYLE_MARKER_RE.match(line)
    if not match:
        return None

    try:
        data = json.loads(base64.b64decode(match.group(1)).decode("utf-8"))
    except Exception:
        return None

    style_id = data.get("style_id")
    style_name = data.get("style_name")
    if not isinstance(style_id, str) or not isinstance(style_name, str):
        return None

    return {"style_id": style_id, "style_name": style_name}


def decode_first_line_indent_marker(line: str) -> Optional[float]:
    match = FIRST_LINE_INDENT_MARKER_RE.match(line)
    if not match:
        return None

    try:
        return float(match.group(1))
    except ValueError:
        return None


def decode_paragraph_alignment_marker(line: str) -> Optional[str]:
    match = PARAGRAPH_ALIGN_MARKER_RE.match(line)
    if not match:
        return None
    return match.group(1)


def decode_table_meta_marker(line: str) -> Optional[Dict[str, Any]]:
    match = TABLE_META_MARKER_RE.match(line)
    if not match:
        return None

    try:
        raw = base64.b64decode(match.group(1))
        try:
            raw = zlib.decompress(raw)
        except zlib.error:
            pass
        table_meta = json.loads(raw.decode("utf-8"))
    except Exception:
        return None

    return table_meta if isinstance(table_meta, dict) else None


def build_paragraph_style_meta(
    style_id: str,
    style_name: str,
    heading_level: Optional[int],
) -> Optional[Dict[str, str]]:
    if not style_id or heading_level is not None:
        return None

    normalized_names = {style_id.strip().lower(), style_name.strip().lower()}
    if "normal" in normalized_names:
        return None

    return {"style_id": style_id, "style_name": style_name}


def find_paragraph_style_by_meta(
    doc: DocxDocument,
    style_meta: Optional[Dict[str, str]],
) -> Optional[ParagraphStyle]:
    if not style_meta:
        return None

    style_id = style_meta.get("style_id", "")
    style_name = style_meta.get("style_name", "")

    for style in doc.styles:
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue
        if style_id and style.style_id == style_id:
            return cast(ParagraphStyle, style)
        if style_name and style.name == style_name:
            return cast(ParagraphStyle, style)

    for candidate in (style_name, style_id):
        if not candidate:
            continue
        try:
            return cast(ParagraphStyle, doc.styles[candidate])
        except Exception:
            pass

    return None


def twips_to_cm(twips: int) -> float:
    return round(twips / TWIPS_PER_CM, 2)


def paragraph_alignment_value(paragraph: etree._Element) -> Optional[str]:
    values = xpath_strings(paragraph, "./w:pPr/w:jc/@w:val")
    return values[0] if values else None


def apply_paragraph_alignment(paragraph, alignment: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    jc = p_pr.find(qn("w:jc"))
    if jc is None:
        jc = etree.Element(qn("w:jc"))
        p_pr.append(jc)
    jc.set(qn("w:val"), alignment)


def paragraph_first_line_indent_cm(paragraph: etree._Element) -> Optional[float]:
    values = xpath_strings(paragraph, "./w:pPr/w:ind/@w:firstLine")
    if not values:
        return None

    try:
        first_line = int(values[0])
    except ValueError:
        return None

    if first_line <= 0:
        return None

    return twips_to_cm(first_line)


def extract_paragraph_meta(
    paragraph: etree._Element,
    style_map: Dict[str, str],
    heading_level: Optional[int] = None,
) -> Dict[str, Any]:
    style_ids = xpath_strings(paragraph, "./w:pPr/w:pStyle/@w:val")
    style_id = style_ids[0] if style_ids else ""
    style_name = style_map.get(style_id, style_id)

    meta: Dict[str, Any] = {}
    style_meta = build_paragraph_style_meta(style_id, style_name, heading_level)
    if style_meta:
        meta["style_meta"] = style_meta

    first_line_indent_cm = paragraph_first_line_indent_cm(paragraph)
    if first_line_indent_cm is not None:
        meta["first_line_indent_cm"] = first_line_indent_cm

    alignment = paragraph_alignment_value(paragraph)
    if alignment and alignment.lower() not in {"left", "start"}:
        meta["alignment"] = alignment

    return meta


def cell_key(row_idx: int, col_idx: int) -> str:
    return f"{row_idx},{col_idx}"


def extract_row_meta(row: etree._Element) -> Dict[str, Any]:
    row_meta: Dict[str, Any] = {}

    height_nodes = xpath_elements(row, "./w:trPr/w:trHeight")
    if height_nodes:
        height_val = height_nodes[0].get(wtag("val"), "")
        height_rule = height_nodes[0].get(wtag("hRule"), "")
        if height_val or height_rule:
            row_meta["height"] = {"val": height_val, "hRule": height_rule}

    if xpath_elements(row, "./w:trPr/w:tblHeader"):
        row_meta["header"] = True

    if xpath_elements(row, "./w:trPr/w:cantSplit"):
        row_meta["cant_split"] = True

    return row_meta


def apply_row_meta(row, row_meta: Optional[Dict[str, Any]]) -> None:
    if not row_meta:
        return

    tr = row._tr
    tr_pr = tr.find(wtag("trPr"))
    if tr_pr is None:
        tr_pr = etree.Element(wtag("trPr"))
        tr.insert(0, tr_pr)

    height_meta = row_meta.get("height")
    if isinstance(height_meta, dict):
        tr_height = tr_pr.find(wtag("trHeight"))
        if tr_height is None:
            tr_height = etree.Element(wtag("trHeight"))
            tr_pr.append(tr_height)
        height_val = height_meta.get("val")
        height_rule = height_meta.get("hRule")
        if isinstance(height_val, str) and height_val:
            tr_height.set(wtag("val"), height_val)
        if isinstance(height_rule, str) and height_rule:
            tr_height.set(wtag("hRule"), height_rule)

    if row_meta.get("header"):
        if tr_pr.find(wtag("tblHeader")) is None:
            tr_pr.append(etree.Element(wtag("tblHeader")))

    if row_meta.get("cant_split"):
        if tr_pr.find(wtag("cantSplit")) is None:
            tr_pr.append(etree.Element(wtag("cantSplit")))


def parse_grid_span(cell: etree._Element) -> int:
    values = xpath_strings(cell, "./w:tcPr/w:gridSpan/@w:val")
    if not values:
        return 1
    try:
        return max(int(values[0]), 1)
    except ValueError:
        return 1


def parse_vmerge_value(cell: etree._Element) -> Optional[str]:
    nodes = xpath_elements(cell, "./w:tcPr/w:vMerge")
    if not nodes:
        return None
    return nodes[0].get(wtag("val")) or "continue"


def extract_cell_meta(cell: etree._Element, style_map: Dict[str, str]) -> Dict[str, Any]:
    cell_meta: Dict[str, Any] = {}

    width_nodes = xpath_elements(cell, "./w:tcPr/w:tcW")
    if width_nodes:
        width_type = width_nodes[0].get(wtag("type"), "")
        width_value = width_nodes[0].get(wtag("w"), "")
        if width_type or width_value:
            cell_meta["width"] = {"type": width_type, "w": width_value}

    vertical_align = xpath_strings(cell, "./w:tcPr/w:vAlign/@w:val")
    if vertical_align:
        cell_meta["vertical_align"] = vertical_align[0]

    shading_fill = xpath_strings(cell, "./w:tcPr/w:shd/@w:fill")
    if shading_fill and shading_fill[0] and shading_fill[0].lower() != "auto":
        cell_meta["shading_fill"] = shading_fill[0]

    paragraph_metas = [extract_paragraph_meta(paragraph, style_map) for paragraph in xpath_elements(cell, "./w:p")]
    if paragraph_metas and (len(paragraph_metas) != 1 or any(meta for meta in paragraph_metas)):
        cell_meta["paragraphs"] = paragraph_metas

    return cell_meta


def table_column_count(table: etree._Element) -> int:
    grid_widths = xpath_elements(table, "./w:tblGrid/w:gridCol")
    if grid_widths:
        return len(grid_widths)

    max_cols = 0
    for row in xpath_elements(table, "./w:tr"):
        col_total = 0
        for cell in xpath_elements(row, "./w:tc"):
            col_total += parse_grid_span(cell)
        max_cols = max(max_cols, col_total)
    return max_cols


def table_cell_to_markdown_text(cell: etree._Element, style_map: Dict[str, str], embed_omml: bool) -> Tuple[str, Dict[str, Any]]:
    paragraph_texts: List[str] = []
    for paragraph in xpath_elements(cell, "./w:p"):
        text = segments_to_inline_markdown(extract_segments_from_container(paragraph), embed_omml)
        paragraph_texts.append(text.replace("|", r"\|").replace("\n", "<br>"))

    visible_paragraphs = list(paragraph_texts)
    while visible_paragraphs and visible_paragraphs[-1] == "":
        visible_paragraphs.pop()

    cell_text = "<br><br>".join(visible_paragraphs)
    return cell_text, extract_cell_meta(cell, style_map)


def collect_table_grid(table: etree._Element, style_map: Dict[str, str], embed_omml: bool) -> Dict[str, Any]:
    row_elements = xpath_elements(table, "./w:tr")
    row_count = len(row_elements)
    col_count = table_column_count(table)
    visible_rows = [["" for _ in range(col_count)] for _ in range(row_count)]
    cell_lookup: Dict[Tuple[int, int], Dict[str, Any]] = {}

    for row_idx, row in enumerate(row_elements):
        col_idx = 0
        for cell in xpath_elements(row, "./w:tc"):
            colspan = parse_grid_span(cell)
            cell_text, cell_meta = table_cell_to_markdown_text(cell, style_map, embed_omml)
            info = {
                "row": row_idx,
                "col": col_idx,
                "colspan": colspan,
                "vmerge": parse_vmerge_value(cell),
                "text": cell_text,
                "cell_meta": cell_meta,
            }
            cell_lookup[(row_idx, col_idx)] = info
            col_idx += colspan

    merges: List[Dict[str, int]] = []
    cells_meta: Dict[str, Any] = {}
    for (row_idx, col_idx), info in cell_lookup.items():
        if info["vmerge"] == "continue":
            continue

        rowspan = 1
        if info["vmerge"] == "restart":
            next_row = row_idx + 1
            while next_row < row_count:
                next_info = cell_lookup.get((next_row, col_idx))
                if next_info is None or next_info["vmerge"] != "continue" or next_info["colspan"] != info["colspan"]:
                    break
                rowspan += 1
                next_row += 1

        visible_rows[row_idx][col_idx] = info["text"]
        if info["cell_meta"]:
            cells_meta[cell_key(row_idx, col_idx)] = info["cell_meta"]

        if info["colspan"] > 1 or rowspan > 1:
            merges.append({"row": row_idx, "col": col_idx, "rowspan": rowspan, "colspan": info["colspan"]})

    return {
        "row_count": row_count,
        "col_count": col_count,
        "rows": visible_rows,
        "merges": merges,
        "cells": cells_meta,
        "rows_meta": [extract_row_meta(row) for row in row_elements],
    }


def extract_table_meta(table: etree._Element, style_map: Dict[str, str]) -> Optional[Dict[str, Any]]:
    table_meta: Dict[str, Any] = {}

    style_ids = xpath_strings(table, "./w:tblPr/w:tblStyle/@w:val")
    if style_ids:
        table_meta["style_id"] = style_ids[0]

    layout_types = xpath_strings(table, "./w:tblPr/w:tblLayout/@w:type")
    if layout_types:
        table_meta["layout"] = layout_types[0]

    alignments = xpath_strings(table, "./w:tblPr/w:jc/@w:val")
    if alignments:
        table_meta["alignment"] = alignments[0]

    table_width_nodes = xpath_elements(table, "./w:tblPr/w:tblW")
    if table_width_nodes:
        table_width_type = table_width_nodes[0].get(wtag("type"), "")
        table_width_value = table_width_nodes[0].get(wtag("w"), "")
        if table_width_type or table_width_value:
            table_meta["table_width"] = {"type": table_width_type, "w": table_width_value}

    grid_widths: List[int] = []
    for grid_col in xpath_elements(table, "./w:tblGrid/w:gridCol"):
        width = grid_col.get(wtag("w"))
        if width is None:
            continue
        try:
            grid_widths.append(int(width))
        except ValueError:
            continue
    if grid_widths:
        table_meta["grid_widths"] = grid_widths

    grid = collect_table_grid(table, style_map, embed_omml=False)
    table_meta["row_count"] = grid["row_count"]
    table_meta["col_count"] = grid["col_count"]
    if grid["merges"]:
        table_meta["merges"] = grid["merges"]
    if grid["cells"]:
        table_meta["cells"] = grid["cells"]
    rows_meta = grid["rows_meta"]
    if rows_meta and any(row_meta for row_meta in rows_meta):
        table_meta["rows_meta"] = rows_meta

    return table_meta or None


def apply_table_meta(table, table_meta: Optional[Dict[str, Any]]) -> None:
    if not table_meta:
        return

    tbl = table._tbl
    tbl_pr = tbl.find(wtag("tblPr"))
    if tbl_pr is None:
        tbl_pr = etree.Element(wtag("tblPr"))
        tbl.insert(0, tbl_pr)

    style_id = table_meta.get("style_id")
    if isinstance(style_id, str) and style_id:
        tbl_style = tbl_pr.find(wtag("tblStyle"))
        if tbl_style is None:
            tbl_style = etree.Element(wtag("tblStyle"))
            tbl_pr.insert(0, tbl_style)
        tbl_style.set(wtag("val"), style_id)

    layout = table_meta.get("layout")
    if isinstance(layout, str) and layout:
        tbl_layout = tbl_pr.find(wtag("tblLayout"))
        if tbl_layout is None:
            tbl_layout = etree.Element(wtag("tblLayout"))
            tbl_pr.append(tbl_layout)
        tbl_layout.set(wtag("type"), layout)
        try:
            table.autofit = layout != "fixed"
        except Exception:
            pass

    alignment = table_meta.get("alignment")
    if isinstance(alignment, str) and alignment:
        tbl_jc = tbl_pr.find(wtag("jc"))
        if tbl_jc is None:
            tbl_jc = etree.Element(wtag("jc"))
            tbl_pr.append(tbl_jc)
        tbl_jc.set(wtag("val"), alignment)

    table_width = table_meta.get("table_width")
    if isinstance(table_width, dict):
        tbl_w = tbl_pr.find(wtag("tblW"))
        if tbl_w is None:
            tbl_w = etree.Element(wtag("tblW"))
            tbl_pr.append(tbl_w)
        width_type = table_width.get("type")
        width_value = table_width.get("w")
        if isinstance(width_type, str) and width_type:
            tbl_w.set(wtag("type"), width_type)
        if isinstance(width_value, str) and width_value:
            tbl_w.set(wtag("w"), width_value)

    grid_widths = table_meta.get("grid_widths")
    if not isinstance(grid_widths, list) or not grid_widths or not all(isinstance(width, int) for width in grid_widths):
        return

    tbl_grid = tbl.find(wtag("tblGrid"))
    if tbl_grid is None:
        tbl_grid = etree.Element(wtag("tblGrid"))
        insert_at = 1 if len(tbl) > 0 and tbl[0].tag == wtag("tblPr") else 0
        tbl.insert(insert_at, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in grid_widths:
        grid_col = etree.Element(wtag("gridCol"))
        grid_col.set(wtag("w"), str(width))
        tbl_grid.append(grid_col)


def heading_level_from_style(style_id: str, style_map: Dict[str, str]) -> Optional[int]:
    candidates = []
    if style_id:
        candidates.append(style_id.lower())
        candidates.append(style_map.get(style_id, "").lower())

    for candidate in candidates:
        match = re.search(r"heading\s*([1-6])", candidate)
        if match:
            return int(match.group(1))
        match_cn = re.search(r"标题\s*([1-6])", candidate)
        if match_cn:
            return int(match_cn.group(1))
    return None


MATH_TEXT_MAP = {
    "−": "-",
    "∈": r"\in ",
    "∉": r"\notin ",
    "≤": r"\leq ",
    "≥": r"\geq ",
    "≠": r"\neq ",
    "≈": r"\approx ",
    "→": r"\to ",
    "←": r"\leftarrow ",
    "↔": r"\leftrightarrow ",
    "⋅": r"\cdot ",
    "·": r"\cdot ",
    "×": r"\times ",
    "∗": r"\ast ",
    "…": r"\ldots ",
}

NARY_MAP = {
    "∑": r"\sum",
    "∏": r"\prod",
    "∫": r"\int",
    "⋃": r"\bigcup",
    "⋂": r"\bigcap",
}

ACCENT_MAP = {
    "̃": r"\tilde",
    "̇": r"\dot",
    "̈": r"\ddot",
    "̂": r"\hat",
    "¯": r"\bar",
}

FUNCTION_NAMES = {
    "sin",
    "cos",
    "tan",
    "exp",
    "log",
    "ln",
    "min",
    "max",
    "lim",
}

LATEX_UNSUPPORTED_SIZING_CMD_RE = re.compile(r"\\(left|right|big|Big|bigg|Bigg)\b")

LATEX_NARY_TO_CHAR = {
    "sum": "∑",
    "prod": "∏",
    "int": "∫",
    "bigcup": "⋃",
    "bigcap": "⋂",
}

LATEX_ACCENT_TO_CHAR = {
    "dot": "̇",
    "ddot": "̈",
    "hat": "̂",
    "tilde": "̃",
    "bar": "¯",
}

LATEX_STYLE_WRAP_COMMANDS = {"mathrm", "mathbf", "mathit", "mathsf", "mathtt", "text"}

LATEX_SCRIPT_STYLE_COMMANDS = {
    "mathcal": "script",
    "mathbb": "double-struck",
}

LATEX_MATRIX_ENVIRONMENTS = {"matrix", "pmatrix", "bmatrix", "vmatrix", "Vmatrix"}

LATEX_MATRIX_DELIMITERS = {
    "pmatrix": ("(", ")"),
    "bmatrix": ("[", "]"),
    "vmatrix": ("|", "|"),
    "Vmatrix": ("‖", "‖"),
}

LATEX_FUNC_COMMANDS = {"sin", "cos", "tan", "exp", "log", "ln", "min", "max", "lim"}

LATEX_COMMAND_TEXT_MAP = {
    "cdot": "·",
    "times": "×",
    "ast": "∗",
    "in": "∈",
    "notin": "∉",
    "leq": "≤",
    "geq": "≥",
    "neq": "≠",
    "approx": "≈",
    "to": "→",
    "leftarrow": "←",
    "leftrightarrow": "↔",
    "ldots": "…",
    "alpha": "α",
    "beta": "β",
    "gamma": "γ",
    "delta": "δ",
    "epsilon": "ϵ",
    "varepsilon": "ε",
    "zeta": "ζ",
    "eta": "η",
    "theta": "θ",
    "vartheta": "ϑ",
    "iota": "ι",
    "kappa": "κ",
    "lambda": "λ",
    "mu": "μ",
    "nu": "ν",
    "xi": "ξ",
    "pi": "π",
    "varpi": "ϖ",
    "rho": "ρ",
    "sigma": "σ",
    "varsigma": "ς",
    "tau": "τ",
    "upsilon": "υ",
    "phi": "ϕ",
    "varphi": "φ",
    "chi": "χ",
    "psi": "ψ",
    "omega": "ω",
    "Gamma": "Γ",
    "Delta": "Δ",
    "Theta": "Θ",
    "Lambda": "Λ",
    "Xi": "Ξ",
    "Pi": "Π",
    "Sigma": "Σ",
    "Upsilon": "Υ",
    "Phi": "Φ",
    "Psi": "Ψ",
    "Omega": "Ω",
}


def normalize_math_text(text: str) -> str:
    return "".join(MATH_TEXT_MAP.get(ch, ch) for ch in text)


def compact_spaces(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\s*\n\s*", " ", text)
    return text.strip()


def group_base(expr: str) -> str:
    expr = expr.strip()
    if not expr:
        return ""
    if re.fullmatch(r"[A-Za-z0-9]", expr):
        return expr
    if re.fullmatch(r"\\[A-Za-z]+", expr):
        return expr
    if expr.startswith("{") and expr.endswith("}"):
        return expr
    return "{" + expr + "}"


def normalize_function_name(name: str) -> str:
    name = compact_spaces(name)
    bare = name.replace("\\", "")
    if name.startswith("\\"):
        return name
    if bare in FUNCTION_NAMES:
        return f"\\{bare}"
    return name


def math_child(node: etree._Element, child_name: str) -> Optional[etree._Element]:
    for child in node:
        if isinstance(child.tag, str) and child.tag == mtag(child_name):
            return child
    return None


def iter_math_children(node: etree._Element) -> Sequence[etree._Element]:
    children: List[etree._Element] = []
    for child in node:
        if not isinstance(child.tag, str):
            continue
        if local_name(child).endswith("Pr"):
            continue
        children.append(child)
    return children


def omml_to_latex(node: Optional[etree._Element]) -> str:
    if node is None:
        return ""

    name = local_name(node)
    if not name:
        return ""

    if name.endswith("Pr"):
        return ""

    if name in {"oMath", "e", "sub", "sup", "num", "den", "fName", "lim"}:
        pieces = [omml_to_latex(child) for child in iter_math_children(node)]
        return compact_spaces("".join(pieces))

    if name == "oMathPara":
        equations = [omml_to_latex(child) for child in node if isinstance(child.tag, str) and child.tag == mtag("oMath")]
        if equations:
            return compact_spaces(" ".join(part for part in equations if part))
        pieces = [omml_to_latex(child) for child in iter_math_children(node)]
        return compact_spaces(" ".join(pieces))

    if name == "r":
        text_parts = []
        for child in node:
            if not isinstance(child.tag, str):
                continue
            if child.tag == mtag("t"):
                text_parts.append(child.text or "")
        if text_parts:
            return normalize_math_text("".join(text_parts))
        pieces = [omml_to_latex(child) for child in iter_math_children(node)]
        return compact_spaces("".join(pieces))

    if name == "t":
        return normalize_math_text(node.text or "")

    if name == "sSub":
        base = group_base(omml_to_latex(math_child(node, "e")))
        sub = omml_to_latex(math_child(node, "sub"))
        return compact_spaces(f"{base}_{{{sub}}}")

    if name == "sSup":
        base = group_base(omml_to_latex(math_child(node, "e")))
        sup = omml_to_latex(math_child(node, "sup"))
        return compact_spaces(f"{base}^{{{sup}}}")

    if name == "sSubSup":
        base = group_base(omml_to_latex(math_child(node, "e")))
        sub = omml_to_latex(math_child(node, "sub"))
        sup = omml_to_latex(math_child(node, "sup"))
        return compact_spaces(f"{base}_{{{sub}}}^{{{sup}}}")

    if name == "f":
        num = omml_to_latex(math_child(node, "num"))
        den = omml_to_latex(math_child(node, "den"))
        return compact_spaces(rf"\frac{{{num}}}{{{den}}}")

    if name == "acc":
        base = omml_to_latex(math_child(node, "e"))
        chr_values = xpath_strings(node, "./m:accPr/m:chr/@m:val")
        accent = chr_values[0] if chr_values else ""
        accent_cmd = ACCENT_MAP.get(accent)
        if accent_cmd:
            return compact_spaces(rf"{accent_cmd}{{{base}}}")
        if accent:
            return compact_spaces(rf"\overset{{{normalize_math_text(accent)}}}{{{base}}}")
        return compact_spaces(rf"\hat{{{base}}}")

    if name == "nary":
        chr_values = xpath_strings(node, "./m:naryPr/m:chr/@m:val")
        symbol = chr_values[0] if chr_values else "∑"
        operator = NARY_MAP.get(symbol, normalize_math_text(symbol))
        sub = omml_to_latex(math_child(node, "sub"))
        sup = omml_to_latex(math_child(node, "sup"))
        expr = omml_to_latex(math_child(node, "e"))

        out = operator
        if sub:
            out += f"_{{{sub}}}"
        if sup:
            out += f"^{{{sup}}}"
        if expr:
            out += f" {expr}"
        return compact_spaces(out)

    if name == "func":
        fname = normalize_function_name(omml_to_latex(math_child(node, "fName")))
        arg = omml_to_latex(math_child(node, "e"))
        if not arg:
            return compact_spaces(fname)
        if arg.startswith("(") and arg.endswith(")"):
            return compact_spaces(f"{fname}{arg}")
        return compact_spaces(rf"{fname}\left({arg}\right)")

    if name == "limLow":
        base = normalize_function_name(omml_to_latex(math_child(node, "e")))
        limit_expr = omml_to_latex(math_child(node, "lim"))
        if limit_expr:
            return compact_spaces(rf"{base}_{{{limit_expr}}}")
        return compact_spaces(base)

    if name == "rad":
        degree = omml_to_latex(math_child(node, "deg"))
        expr = omml_to_latex(math_child(node, "e"))
        if degree:
            return compact_spaces(rf"\sqrt[{degree}]{{{expr}}}")
        return compact_spaces(rf"\sqrt{{{expr}}}")

    if name == "lim":
        pieces = [omml_to_latex(child) for child in iter_math_children(node)]
        return compact_spaces("".join(pieces))

    pieces = [omml_to_latex(child) for child in iter_math_children(node)]
    if pieces:
        return compact_spaces("".join(pieces))

    text_fallback = "".join(xpath_strings(node, ".//m:t/text()"))
    return compact_spaces(normalize_math_text(text_fallback))


def sanitize_latex_for_word(latex: str) -> str:
    text = latex.replace("\u2009", " ").replace("\u00a0", " ")
    text = LATEX_UNSUPPORTED_SIZING_CMD_RE.sub("", text)
    text = text.replace(r"\,", " ").replace(r"\;", " ").replace(r"\!", "")
    text = re.sub(r"\\(quad|qquad)\b", " ", text)
    text = text.replace(r"\(", "(").replace(r"\)", ")")
    text = text.replace(r"\{", "{").replace(r"\}", "}")
    text = text.replace("\n", " ")
    return compact_spaces(text)


def _xml_attr_escape(text: str) -> str:
    return xml_escape(text, {'"': "&quot;"})


class LatexParser:
    def __init__(self, text: str):
        self.tokens = self._tokenize(text)
        self.pos = 0

    @staticmethod
    def _tokenize(text: str) -> List[Tuple[str, str]]:
        tokens: List[Tuple[str, str]] = []
        i = 0
        while i < len(text):
            ch = text[i]

            if ch.isspace():
                i += 1
                continue

            if ch == "\\":
                if i + 1 < len(text) and text[i + 1].isalpha():
                    j = i + 2
                    while j < len(text) and text[j].isalpha():
                        j += 1
                    tokens.append(("cmd", text[i + 1 : j]))
                    i = j
                    continue

                if i + 1 < len(text):
                    tokens.append(("text", text[i + 1]))
                    i += 2
                    continue

                i += 1
                continue

            if ch == "{":
                tokens.append(("lbrace", ch))
                i += 1
                continue
            if ch == "}":
                tokens.append(("rbrace", ch))
                i += 1
                continue
            if ch == "[":
                tokens.append(("lbracket", ch))
                i += 1
                continue
            if ch == "]":
                tokens.append(("rbracket", ch))
                i += 1
                continue
            if ch == "_":
                tokens.append(("underscore", ch))
                i += 1
                continue
            if ch == "^":
                tokens.append(("caret", ch))
                i += 1
                continue

            tokens.append(("text", ch))
            i += 1

        return tokens

    def _peek(self) -> Tuple[str, str]:
        if self.pos >= len(self.tokens):
            return ("eof", "")
        return self.tokens[self.pos]

    def _accept(self, token_type: str) -> bool:
        if self._peek()[0] == token_type:
            self.pos += 1
            return True
        return False

    def _advance(self) -> Tuple[str, str]:
        token = self._peek()
        if token[0] != "eof":
            self.pos += 1
        return token

    def parse(self) -> Any:
        return self._parse_expression({"eof"})

    def _parse_expression(self, stop_types: set[str]) -> Any:
        items: List[Any] = []

        while True:
            token_type = self._peek()[0]
            if token_type in stop_types or token_type == "eof":
                break

            base = self._parse_primary()
            if base is None:
                break

            sub = None
            sup = None
            while True:
                token_type = self._peek()[0]
                if token_type == "underscore":
                    self._advance()
                    sub = self._parse_script_operand()
                    continue
                if token_type == "caret":
                    self._advance()
                    sup = self._parse_script_operand()
                    continue
                break

            if sub is not None or sup is not None:
                base = _attach_scripts(base, sub, sup)

            items.append(base)

        items = _fuse_nary(items)
        if not items:
            return ("empty",)
        if len(items) == 1:
            return items[0]
        return ("seq", tuple(items))

    def _parse_primary(self) -> Optional[Any]:
        token_type, token_val = self._peek()

        if token_type in {"rbrace", "rbracket", "eof"}:
            return None

        if token_type == "lbrace":
            self._advance()
            node = self._parse_expression({"rbrace", "eof"})
            self._accept("rbrace")
            return node

        if token_type == "lbracket":
            self._advance()
            node = self._parse_expression({"rbracket", "eof"})
            self._accept("rbracket")
            return node

        if token_type == "cmd":
            self._advance()
            return self._parse_command(token_val)

        self._advance()
        return ("text", token_val)

    def _parse_group_or_primary(self) -> Any:
        token_type = self._peek()[0]
        if token_type == "lbrace":
            self._advance()
            node = self._parse_expression({"rbrace", "eof"})
            self._accept("rbrace")
            return node
        return self._parse_primary() or ("empty",)

    def _parse_environment_name(self) -> str:
        if not self._accept("lbrace"):
            return ""

        parts: List[str] = []
        while True:
            token_type, token_val = self._peek()
            if token_type in {"rbrace", "eof"}:
                break
            self._advance()
            if token_type in {"text", "cmd"}:
                parts.append(token_val)

        self._accept("rbrace")
        return "".join(parts).strip()

    def _peek_is_environment_end(self, env_name: str) -> bool:
        token_type, token_val = self._peek()
        if token_type != "cmd" or token_val != "end":
            return False

        save_pos = self.pos
        self._advance()
        parsed_env = self._parse_environment_name()
        self.pos = save_pos
        return parsed_env == env_name

    def _parse_expression_until_matrix_separator(self, env_name: str) -> Any:
        items: List[Any] = []

        while True:
            token_type, token_val = self._peek()
            if token_type == "eof" or self._peek_is_environment_end(env_name):
                break
            if token_type == "text" and token_val in {"&", "\\"}:
                break

            start_pos = self.pos
            base = self._parse_primary()
            if base is None:
                break

            sub = None
            sup = None
            while True:
                token_type = self._peek()[0]
                if token_type == "underscore":
                    self._advance()
                    sub = self._parse_script_operand()
                    continue
                if token_type == "caret":
                    self._advance()
                    sup = self._parse_script_operand()
                    continue
                break

            if sub is not None or sup is not None:
                base = _attach_scripts(base, sub, sup)

            items.append(base)

            if self.pos == start_pos:
                self._advance()

        items = _fuse_nary(items)
        if not items:
            return ("empty",)
        if len(items) == 1:
            return items[0]
        return ("seq", tuple(items))

    def _parse_matrix_environment(self, env_name: str) -> Any:
        rows: List[List[Any]] = []
        current_row: List[Any] = []

        while True:
            if self._peek()[0] == "eof":
                if current_row or not rows:
                    rows.append(current_row if current_row else [("empty",)])
                break

            if self._peek_is_environment_end(env_name):
                if current_row or not rows:
                    rows.append(current_row if current_row else [("empty",)])
                self._advance()
                self._parse_environment_name()
                break

            start_pos = self.pos
            cell_expr = self._parse_expression_until_matrix_separator(env_name)
            current_row.append(cell_expr)

            token_type, token_val = self._peek()
            if token_type == "text" and token_val == "&":
                self._advance()
                continue

            if token_type == "text" and token_val == "\\":
                self._advance()
                rows.append(current_row)
                current_row = []
                continue

            if self._peek_is_environment_end(env_name):
                continue

            if self.pos == start_pos:
                self._advance()

        normalized_rows: List[Tuple[Any, ...]] = []
        for row in rows:
            if not row:
                normalized_rows.append((("empty",),))
            else:
                normalized_rows.append(tuple(row))

        if not normalized_rows:
            normalized_rows = [(("empty",),)]

        return ("matrix", env_name, tuple(normalized_rows))

    def _parse_script_operand(self) -> Any:
        return self._parse_group_or_primary()

    def _parse_command(self, name: str) -> Any:
        if name == "begin":
            env_name = self._parse_environment_name()
            if env_name in LATEX_MATRIX_ENVIRONMENTS:
                return self._parse_matrix_environment(env_name)
            return ("text", env_name)

        if name == "end":
            self._parse_environment_name()
            return ("empty",)

        if name == "frac":
            num = self._parse_group_or_primary()
            den = self._parse_group_or_primary()
            return ("frac", num, den)

        if name == "sqrt":
            degree = None
            if self._accept("lbracket"):
                degree = self._parse_expression({"rbracket", "eof"})
                self._accept("rbracket")
            base = self._parse_group_or_primary()
            return ("sqrt", degree, base)

        if name in LATEX_ACCENT_TO_CHAR:
            base = self._parse_group_or_primary()
            return ("acc", LATEX_ACCENT_TO_CHAR[name], base)

        if name in LATEX_NARY_TO_CHAR:
            return ("nary", LATEX_NARY_TO_CHAR[name], None, None, None)

        if name in LATEX_SCRIPT_STYLE_COMMANDS:
            return ("style", LATEX_SCRIPT_STYLE_COMMANDS[name], self._parse_group_or_primary())

        if name in LATEX_STYLE_WRAP_COMMANDS:
            return self._parse_group_or_primary()

        if name in LATEX_FUNC_COMMANDS:
            return ("text", name)

        mapped = LATEX_COMMAND_TEXT_MAP.get(name)
        if mapped is not None:
            return ("text", mapped)

        return ("text", name)


def _attach_scripts(base: Any, sub: Any, sup: Any) -> Any:
    if isinstance(base, tuple) and base and base[0] == "nary":
        return ("nary", base[1], sub if sub is not None else base[2], sup if sup is not None else base[3], base[4])

    if sub is not None and sup is not None:
        return ("subsup", base, sub, sup)
    if sub is not None:
        return ("sub", base, sub)
    if sup is not None:
        return ("sup", base, sup)
    return base


def _fuse_nary(items: List[Any]) -> List[Any]:
    fused: List[Any] = []
    i = 0
    while i < len(items):
        node = items[i]
        if isinstance(node, tuple) and node and node[0] == "nary" and node[4] is None and i + 1 < len(items):
            nxt = items[i + 1]
            if not (isinstance(nxt, tuple) and nxt and nxt[0] == "text" and nxt[1] in {",", ";", ":"}):
                node = ("nary", node[1], node[2], node[3], nxt)
                i += 1
        fused.append(node)
        i += 1
    return fused


def _omml_text_run(text: str) -> str:
    if not text:
        return ""
    return f"<m:r><m:t>{xml_escape(text)}</m:t></m:r>"


def apply_math_script_to_omml(omml: str, script_value: str) -> str:
    if not omml:
        return omml

    wrapper = f'<m:root xmlns:m="{M_NS}">{omml}</m:root>'
    root = etree.fromstring(wrapper.encode("utf-8"))

    for run in xpath_elements(root, ".//m:r"):
        r_pr = run.find(mtag("rPr"))
        if r_pr is None:
            r_pr = etree.Element(mtag("rPr"))
            run.insert(0, r_pr)

        for child in list(r_pr):
            if child.tag == mtag("scr"):
                r_pr.remove(child)

        scr = etree.Element(mtag("scr"))
        scr.set(f"{{{M_NS}}}val", script_value)
        r_pr.insert(0, scr)

    return "".join(etree.tostring(child, encoding="unicode") for child in root)


def latex_node_to_omml(node: Any) -> str:
    if node is None:
        return ""

    if not isinstance(node, tuple) or not node:
        return _omml_text_run(str(node))

    kind = node[0]

    if kind == "empty":
        return ""

    if kind == "text":
        return _omml_text_run(node[1])

    if kind == "seq":
        return "".join(latex_node_to_omml(child) for child in node[1])

    if kind == "style":
        script_style = node[1]
        inner = latex_node_to_omml(node[2])
        return apply_math_script_to_omml(inner, script_style)

    if kind == "frac":
        num = latex_node_to_omml(node[1])
        den = latex_node_to_omml(node[2])
        return f"<m:f><m:num>{num}</m:num><m:den>{den}</m:den></m:f>"

    if kind == "sub":
        base = latex_node_to_omml(node[1])
        sub = latex_node_to_omml(node[2])
        return f"<m:sSub><m:e>{base}</m:e><m:sub>{sub}</m:sub></m:sSub>"

    if kind == "sup":
        base = latex_node_to_omml(node[1])
        sup = latex_node_to_omml(node[2])
        return f"<m:sSup><m:e>{base}</m:e><m:sup>{sup}</m:sup></m:sSup>"

    if kind == "subsup":
        base = latex_node_to_omml(node[1])
        sub = latex_node_to_omml(node[2])
        sup = latex_node_to_omml(node[3])
        return f"<m:sSubSup><m:e>{base}</m:e><m:sub>{sub}</m:sub><m:sup>{sup}</m:sup></m:sSubSup>"

    if kind == "sqrt":
        degree = latex_node_to_omml(node[1]) if node[1] is not None else ""
        base = latex_node_to_omml(node[2])
        deg_xml = f"<m:deg>{degree}</m:deg>" if degree else ""
        return f"<m:rad>{deg_xml}<m:e>{base}</m:e></m:rad>"

    if kind == "acc":
        accent = _xml_attr_escape(node[1])
        base = latex_node_to_omml(node[2])
        return f"<m:acc><m:accPr><m:chr m:val=\"{accent}\"/></m:accPr><m:e>{base}</m:e></m:acc>"

    if kind == "nary":
        symbol = _xml_attr_escape(node[1])
        sub = latex_node_to_omml(node[2]) if node[2] is not None else ""
        sup = latex_node_to_omml(node[3]) if node[3] is not None else ""
        expr = latex_node_to_omml(node[4]) if node[4] is not None else ""
        sub_xml = f"<m:sub>{sub}</m:sub>" if sub else ""
        sup_xml = f"<m:sup>{sup}</m:sup>" if sup else ""
        return (
            f"<m:nary><m:naryPr><m:chr m:val=\"{symbol}\"/></m:naryPr>"
            f"{sub_xml}{sup_xml}<m:e>{expr}</m:e></m:nary>"
        )

    if kind == "matrix":
        env_name = node[1]
        rows = node[2]
        row_xml_parts: List[str] = []
        for row in rows:
            cells_xml = "".join(f"<m:e>{latex_node_to_omml(cell)}</m:e>" for cell in row)
            row_xml_parts.append(f"<m:mr>{cells_xml}</m:mr>")

        matrix_core = f"<m:m>{''.join(row_xml_parts)}</m:m>"
        delimiters = LATEX_MATRIX_DELIMITERS.get(env_name)
        if delimiters is None:
            return matrix_core

        beg, end = delimiters
        return (
            f"<m:d><m:dPr><m:begChr m:val=\"{_xml_attr_escape(beg)}\"/>"
            f"<m:endChr m:val=\"{_xml_attr_escape(end)}\"/></m:dPr>"
            f"<m:e>{matrix_core}</m:e></m:d>"
        )

    return _omml_text_run(str(node))


def latex_to_omml_element(latex: str, block: bool = False) -> Optional[etree._Element]:
    normalized = sanitize_latex_for_word(latex)
    if not normalized:
        return None

    parser = LatexParser(normalized)
    node = parser.parse()
    inner = latex_node_to_omml(node)
    if not inner:
        return None

    if block:
        xml = f'<m:oMathPara xmlns:m="{M_NS}"><m:oMath>{inner}</m:oMath></m:oMathPara>'
    else:
        xml = f'<m:oMath xmlns:m="{M_NS}">{inner}</m:oMath>'

    try:
        return parse_xml(xml)
    except Exception:
        return None


def append_latex_math_to_paragraph(paragraph, latex: str, block: bool = False) -> bool:
    element = latex_to_omml_element(latex, block=block)
    if element is None:
        return False
    paragraph._p.append(element)
    return True


def encode_omml_payload(node: etree._Element) -> str:
    raw = etree.tostring(node, encoding="utf-8", with_tail=False)
    packed = zlib.compress(raw, level=9)
    return base64.b64encode(packed).decode("ascii")


def decode_omml_payload(payload: str) -> Optional[bytes]:
    try:
        packed = base64.b64decode(payload.encode("ascii"))
        return zlib.decompress(packed)
    except Exception:
        return None


def run_to_text(run: etree._Element) -> str:
    parts: List[str] = []
    for child in run:
        if not isinstance(child.tag, str):
            continue
        tag = child.tag
        if tag == wtag("t"):
            parts.append(child.text or "")
        elif tag == wtag("tab"):
            parts.append("\t")
        elif tag in {wtag("br"), wtag("cr")}:
            parts.append("\n")
        elif tag == wtag("noBreakHyphen"):
            parts.append("-")
    return "".join(parts)


def is_run_italic(run: etree._Element) -> bool:
    italic_nodes = xpath_elements(run, "./w:rPr/w:i")
    if not italic_nodes:
        return False
    val = italic_nodes[0].get(wtag("val"))
    if val is None:
        return True
    return val.lower() not in {"0", "false", "off"}


def append_text_token(
    segments: List[Segment],
    text: str,
    bold: bool = False,
    italic: bool = False,
    color: Optional[str] = None,
    underline: Optional[str] = None,
    strike: bool = False,
    background_color: Optional[str] = None,
    highlight: Optional[str] = None,
    character_style: Optional[str] = None,
) -> None:
    if not text:
        return

    normalized_color = normalize_color_value(color)
    normalized_underline = normalize_underline_value(underline)
    normalized_background = normalize_color_value(background_color)
    normalized_highlight = normalize_highlight_value(highlight)
    normalized_character_style = character_style.strip() if character_style else None
    if normalized_character_style == "":
        normalized_character_style = None

    if segments and isinstance(segments[-1], TextToken):
        previous = segments[-1]
        if (
            previous.bold == bold
            and previous.italic == italic
            and previous.color == normalized_color
            and previous.underline == normalized_underline
            and previous.strike == strike
            and previous.background_color == normalized_background
            and previous.highlight == normalized_highlight
            and previous.character_style == normalized_character_style
        ):
            previous.text += text
            return

    segments.append(
        TextToken(
            text=text,
            bold=bold,
            italic=italic,
            color=normalized_color,
            underline=normalized_underline,
            strike=strike,
            background_color=normalized_background,
            highlight=normalized_highlight,
            character_style=normalized_character_style,
        )
    )


def render_text_token(token: TextToken) -> str:
    marker = ""
    if token.bold and token.italic:
        marker = "***"
    elif token.bold:
        marker = "**"
    elif token.italic:
        marker = "*"

    content = token.text if not marker else f"{marker}{token.text}{marker}"
    if token.strike:
        content = f"~~{content}~~"
    if token.underline:
        underline_attr = ""
        if token.underline != "single":
            underline_attr = f' data-docx-underline="{xml_escape(token.underline)}"'
        content = f"<u{underline_attr}>{content}</u>"

    span_attrs: List[str] = []
    style_parts: List[str] = []
    if token.color:
        style_parts.append(f"color:#{token.color}")
    if token.background_color:
        style_parts.append(f"background-color:#{token.background_color}")
    if style_parts:
        span_attrs.append(f'style="{";".join(style_parts)}"')
    if token.highlight:
        span_attrs.append(f'data-docx-highlight="{xml_escape(token.highlight)}"')
    if token.character_style:
        span_attrs.append(f'data-docx-rstyle="{xml_escape(token.character_style)}"')

    if span_attrs:
        return f"<span {' '.join(span_attrs)}>{content}</span>"
    return content


def segments_to_inline_markdown(segments: List[Segment], embed_omml: bool) -> str:
    parts: List[str] = []
    for segment in segments:
        if isinstance(segment, TextToken):
            parts.append(render_text_token(segment))
            continue

        parts.append(f"${segment.latex.strip()}$")
        if embed_omml and segment.omml_payload:
            parts.append(f"<!--OMML_INLINE_Z:{segment.omml_payload}-->")

    return "".join(parts)


def extract_segments_from_container(container: etree._Element) -> List[Segment]:
    segments: List[Segment] = []

    for child in container:
        if not isinstance(child.tag, str):
            continue

        tag = child.tag

        if tag == wtag("r"):
            text = run_to_text(child)
            if text:
                append_text_token(
                    segments,
                    text,
                    bold=is_run_bold(child),
                    italic=is_run_italic(child),
                    color=run_color_value(child),
                    underline=run_underline_value(child),
                    strike=run_strike_value(child),
                    background_color=run_background_color_value(child),
                    highlight=run_highlight_value(child),
                    character_style=run_character_style_value(child),
                )
            continue

        if tag in {wtag("hyperlink"), wtag("smartTag"), wtag("sdt"), wtag("ins")}:
            segments.extend(extract_segments_from_container(child))
            continue

        if tag == mtag("oMath"):
            latex = omml_to_latex(child)
            payload = encode_omml_payload(child)
            segments.append(MathToken(latex=latex, omml_payload=payload, block=False))
            continue

        if tag == mtag("oMathPara"):
            latex = omml_to_latex(child)
            payload = encode_omml_payload(child)
            segments.append(MathToken(latex=latex, omml_payload=payload, block=True))
            continue

    return segments


def paragraph_to_markdown(
    segments: List[Segment],
    heading_level: Optional[int],
    embed_omml: bool,
    style_meta: Optional[Dict[str, str]] = None,
    first_line_indent_cm: Optional[float] = None,
    alignment: Optional[str] = None,
) -> List[str]:
    non_empty_text = "".join(seg.text for seg in segments if isinstance(seg, TextToken)).strip()
    math_segments = [seg for seg in segments if isinstance(seg, MathToken)]
    marker_lines = []
    if style_meta:
        marker_lines.append(encode_paragraph_style_marker(style_meta["style_id"], style_meta["style_name"]))
    if first_line_indent_cm is not None:
        marker_lines.append(encode_first_line_indent_marker(first_line_indent_cm))
    if alignment:
        marker_lines.append(encode_paragraph_alignment_marker(alignment))

    if len(math_segments) == 1 and not non_empty_text and math_segments[0].block:
        token = math_segments[0]
        lines = ["$$", token.latex.strip(), "$$"]
        if embed_omml and token.omml_payload:
            lines.append(f"<!--OMML_BLOCK_Z:{token.omml_payload}-->")
        return marker_lines + lines

    content = segments_to_inline_markdown(segments, embed_omml).strip()
    if heading_level is not None:
        prefix = "#" * heading_level
        return [f"{prefix} {content}".rstrip()]

    if content:
        return marker_lines + [content]
    return marker_lines + [EMPTY_PARAGRAPH_MARKER]


def docx_to_markdown(input_docx: Path, output_md: Path, embed_omml: bool = True) -> None:
    document_root, styles_root = read_xml_from_docx(input_docx)
    style_map = style_id_to_name(styles_root)

    body = xpath_elements(document_root, "/w:document/w:body")
    if not body:
        raise RuntimeError("Invalid DOCX: missing document body")

    lines: List[str] = []
    for child in body[0]:
        if not isinstance(child.tag, str):
            continue

        if child.tag == wtag("p"):
            style_ids = xpath_strings(child, "./w:pPr/w:pStyle/@w:val")
            style_id = style_ids[0] if style_ids else ""
            style_name = style_map.get(style_id, style_id)
            heading_level = heading_level_from_style(style_id, style_map)
            style_meta = build_paragraph_style_meta(style_id, style_name, heading_level)
            first_line_indent_cm = paragraph_first_line_indent_cm(child)
            alignment = paragraph_alignment_value(child)
            if alignment and alignment.lower() in {"left", "start"}:
                alignment = None
            segments = extract_segments_from_container(child)
            para_lines = paragraph_to_markdown(
                segments,
                heading_level,
                embed_omml,
                style_meta,
                first_line_indent_cm,
                alignment,
            )
            lines.extend(para_lines)
            lines.append("")
            continue

        if child.tag == wtag("tbl"):
            lines.extend(table_to_markdown_lines(child, style_map, embed_omml))
            lines.append("")

    while lines and not lines[-1].strip():
        lines.pop()

    output_md.write_text("\n".join(lines) + "\n", encoding="utf-8")


def find_unescaped_dollar(text: str, start: int) -> int:
    i = start
    while i < len(text):
        if text[i] == "$" and (i == 0 or text[i - 1] != "\\"):
            return i
        i += 1
    return -1


def split_inline_segments(text: str) -> List[MarkdownInlineSegment]:
    segments: List[MarkdownInlineSegment] = []
    i = 0

    while i < len(text):
        start = find_unescaped_dollar(text, i)
        if start < 0:
            tail = INLINE_MARKER_RE.sub("", text[i:])
            if tail:
                segments.append(tail)
            break

        leading = INLINE_MARKER_RE.sub("", text[i:start])
        if leading:
            segments.append(leading)

        end = find_unescaped_dollar(text, start + 1)
        if end < 0:
            tail = INLINE_MARKER_RE.sub("", text[start:])
            if tail:
                segments.append(tail)
            break

        latex = text[start + 1 : end]
        marker_match = INLINE_MARKER_RE.match(text, end + 1)
        payload = marker_match.group(1) if marker_match else None
        next_index = marker_match.end() if marker_match else (end + 1)

        segments.append(MathToken(latex=latex, omml_payload=payload, block=False))
        i = next_index

    return segments


def split_emphasis_spans(text: str) -> List[Tuple[str, bool, bool]]:
    spans: List[Tuple[str, bool, bool]] = []
    pattern = re.compile(r"\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*", re.DOTALL)
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            spans.append((text[pos : match.start()], False, False))
        if match.group(1) is not None:
            spans.append((match.group(1), True, True))
        elif match.group(2) is not None:
            spans.append((match.group(2), True, False))
        else:
            spans.append((match.group(3), False, True))
        pos = match.end()
    if pos < len(text):
        spans.append((text[pos:], False, False))
    if not spans:
        spans.append((text, False, False))
    return spans


def split_markdown_table_cells(line: str) -> List[str]:
    raw = line.strip()
    if raw.startswith("|"):
        raw = raw[1:]
    if raw.endswith("|"):
        raw = raw[:-1]

    cells: List[str] = []
    current: List[str] = []
    escaped = False

    for ch in raw:
        if escaped:
            current.append(ch)
            escaped = False
            continue

        if ch == "\\":
            current.append(ch)
            escaped = True
            continue

        if ch == "|":
            cells.append("".join(current).strip())
            current = []
            continue

        current.append(ch)

    cells.append("".join(current).strip())
    return [cell.replace(r"\|", "|") for cell in cells]


def looks_like_markdown_table_row(line: str) -> bool:
    stripped = line.strip()
    if not stripped or "|" not in stripped:
        return False
    return len(split_markdown_table_cells(stripped)) >= 2


def is_markdown_table_start(lines: List[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    return looks_like_markdown_table_row(lines[index]) and bool(TABLE_SEPARATOR_RE.match(lines[index + 1]))


def normalize_markdown_table_rows(rows: List[List[str]]) -> List[List[str]]:
    if not rows:
        return rows
    max_cols = max(len(row) for row in rows)
    if max_cols < 2:
        max_cols = 2
    return [row + [""] * (max_cols - len(row)) for row in rows]


def is_run_bold(run: etree._Element) -> bool:
    bold_nodes = xpath_elements(run, "./w:rPr/w:b")
    if not bold_nodes:
        return False
    val = bold_nodes[0].get(wtag("val"))
    if val is None:
        return True
    return val.lower() not in {"0", "false", "off"}


def normalize_br_tags(text: str) -> str:
    return BR_TAG_RE.sub("<br>", text)


def split_markdown_cell_paragraphs(cell_text: str) -> List[str]:
    normalized = normalize_br_tags(cell_text)
    if not normalized:
        return [""]
    return DOUBLE_BR_TAG_RE.split(normalized)


def build_covered_positions_from_merges(merges: Sequence[Dict[str, int]]) -> set[Tuple[int, int]]:
    covered: set[Tuple[int, int]] = set()
    for merge in merges:
        row = merge["row"]
        col = merge["col"]
        rowspan = merge["rowspan"]
        colspan = merge["colspan"]
        for row_idx in range(row, row + rowspan):
            for col_idx in range(col, col + colspan):
                if row_idx == row and col_idx == col:
                    continue
                covered.add((row_idx, col_idx))
    return covered


def clear_cell_content(cell) -> None:
    tc = cell._tc
    for child in list(tc):
        if child.tag != wtag("tcPr"):
            tc.remove(child)


def apply_cell_meta(cell, cell_meta: Optional[Dict[str, Any]]) -> None:
    if not cell_meta:
        return

    tc = cell._tc
    tc_pr = tc.find(wtag("tcPr"))
    if tc_pr is None:
        tc_pr = etree.Element(wtag("tcPr"))
        tc.insert(0, tc_pr)

    width_meta = cell_meta.get("width")
    if isinstance(width_meta, dict):
        tc_w = tc_pr.find(wtag("tcW"))
        if tc_w is None:
            tc_w = etree.Element(wtag("tcW"))
            tc_pr.insert(0, tc_w)
        width_type = width_meta.get("type")
        width_value = width_meta.get("w")
        if isinstance(width_type, str) and width_type:
            tc_w.set(wtag("type"), width_type)
        if isinstance(width_value, str) and width_value:
            tc_w.set(wtag("w"), width_value)

    vertical_align = cell_meta.get("vertical_align")
    if isinstance(vertical_align, str) and vertical_align:
        v_align = tc_pr.find(wtag("vAlign"))
        if v_align is None:
            v_align = etree.Element(wtag("vAlign"))
            tc_pr.append(v_align)
        v_align.set(wtag("val"), vertical_align)

    shading_fill = cell_meta.get("shading_fill")
    if isinstance(shading_fill, str) and shading_fill:
        shd = tc_pr.find(wtag("shd"))
        if shd is None:
            shd = etree.Element(wtag("shd"))
            tc_pr.append(shd)
        shd.set(wtag("val"), "clear")
        shd.set(wtag("color"), "auto")
        shd.set(wtag("fill"), shading_fill)


def table_to_markdown_lines(table: etree._Element, style_map: Dict[str, str], embed_omml: bool) -> List[str]:
    grid = collect_table_grid(table, style_map, embed_omml)
    rows = cast(List[List[str]], grid["rows"])
    rows = normalize_markdown_table_rows(rows)
    if not rows:
        return ["|  |", "| --- |"]

    column_count = len(rows[0])
    separator = ["---"] * column_count

    lines: List[str] = []
    table_meta = extract_table_meta(table, style_map)
    if table_meta:
        lines.append(encode_table_meta_marker(table_meta))

    lines.extend([f"| {' | '.join(rows[0])} |", f"| {' | '.join(separator)} |"])
    for row in rows[1:]:
        lines.append(f"| {' | '.join(row)} |")
    return lines


def apply_first_line_indent(paragraph, indent_cm: float = DEFAULT_FIRST_LINE_INDENT_CM) -> None:
    paragraph.paragraph_format.first_line_indent = Cm(indent_cm)


def append_markdown_inline_content(
    paragraph,
    text: str,
    prefer_omml: bool,
    doc: Optional[DocxDocument] = None,
    default_run_font: Optional[Dict[str, str]] = None,
) -> None:
    for styled_text, inline_style in split_style_spans(text):
        for underline_text, underline_value in split_underline_spans(styled_text):
            effective_underline = underline_value or inline_style.underline
            for strike_text, has_strike in split_strike_spans(underline_text):
                effective_strike = has_strike or inline_style.strike
                for span_text, is_bold, is_italic in split_emphasis_spans(strike_text):
                    segments = split_inline_segments(span_text)
                    for segment in segments:
                        if isinstance(segment, str):
                            if segment:
                                run = paragraph.add_run(segment)
                                apply_character_style_to_run(run, doc, inline_style.character_style)
                                apply_font_spec_to_run(run, default_run_font)
                                apply_text_color_to_run(run, inline_style.color)
                                apply_run_background_to_run(
                                    run,
                                    background_color=inline_style.background_color,
                                    highlight=inline_style.highlight,
                                )
                                apply_run_underline_to_run(run, effective_underline)
                                apply_run_strike_to_run(run, effective_strike)
                                if is_bold:
                                    run.bold = True
                                if is_italic:
                                    run.italic = True
                            continue

                        restored = False
                        if prefer_omml and segment.omml_payload:
                            restored = append_omml_to_paragraph(paragraph, segment.omml_payload)
                        if not restored:
                            restored = append_latex_math_to_paragraph(paragraph, segment.latex, block=False)
                        if not restored:
                            run = paragraph.add_run(f"${segment.latex}$")
                            apply_character_style_to_run(run, doc, inline_style.character_style)
                            apply_font_spec_to_run(run, default_run_font)
                            apply_text_color_to_run(run, inline_style.color)
                            apply_run_background_to_run(
                                run,
                                background_color=inline_style.background_color,
                                highlight=inline_style.highlight,
                            )
                            apply_run_underline_to_run(run, effective_underline)
                            apply_run_strike_to_run(run, effective_strike)
                            if is_bold:
                                run.bold = True
                            if is_italic:
                                run.italic = True


def append_markdown_cell_content(
    cell,
    cell_text: str,
    prefer_omml: bool,
    doc: DocxDocument,
    default_run_font: Optional[Dict[str, str]] = None,
    cell_meta: Optional[Dict[str, Any]] = None,
) -> None:
    clear_cell_content(cell)

    paragraph_texts = split_markdown_cell_paragraphs(cell_text)
    paragraph_metas: List[Dict[str, Any]] = []
    if isinstance(cell_meta, dict):
        raw_paragraphs = cell_meta.get("paragraphs")
        if isinstance(raw_paragraphs, list):
            paragraph_metas = [meta if isinstance(meta, dict) else {} for meta in raw_paragraphs]

    paragraph_count = max(len(paragraph_texts), len(paragraph_metas), 1)
    for idx in range(paragraph_count):
        paragraph = cell.add_paragraph()
        paragraph_meta = paragraph_metas[idx] if idx < len(paragraph_metas) else {}
        if isinstance(paragraph_meta, dict):
            paragraph_style = find_paragraph_style_by_meta(doc, paragraph_meta.get("style_meta"))
            if paragraph_style is not None:
                paragraph.style = paragraph_style
            alignment = paragraph_meta.get("alignment")
            if isinstance(alignment, str) and alignment:
                apply_paragraph_alignment(paragraph, alignment)
            indent_cm = paragraph_meta.get("first_line_indent_cm")
            if isinstance(indent_cm, (int, float)):
                apply_first_line_indent(paragraph, float(indent_cm))

        paragraph_text = paragraph_texts[idx] if idx < len(paragraph_texts) else ""
        parts = normalize_br_tags(paragraph_text).split("<br>")
        for part_idx, part in enumerate(parts):
            if part_idx > 0:
                break_run = paragraph.add_run()
                apply_font_spec_to_run(break_run, default_run_font)
                break_run.add_break()
            append_markdown_inline_content(paragraph, part, prefer_omml, doc, default_run_font)


def clear_paragraph_content(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag != wtag("pPr"):
            p.remove(child)


def parse_markdown_blocks(markdown_text: str) -> List[dict]:
    lines = markdown_text.splitlines()
    blocks: List[dict] = []
    heading_re = re.compile(r"^(#{1,6})\s+(.*)$")
    pending_paragraph_style: Optional[Dict[str, str]] = None
    pending_first_line_indent_cm: Optional[float] = None
    pending_alignment: Optional[str] = None
    pending_table_meta: Optional[Dict[str, Any]] = None

    i = 0
    while i < len(lines):
        line = lines[i]

        paragraph_style_meta = decode_paragraph_style_marker(line)
        if paragraph_style_meta is not None:
            pending_paragraph_style = paragraph_style_meta
            i += 1
            continue

        first_line_indent_cm = decode_first_line_indent_marker(line)
        if first_line_indent_cm is not None:
            pending_first_line_indent_cm = first_line_indent_cm
            i += 1
            continue

        alignment = decode_paragraph_alignment_marker(line)
        if alignment is not None:
            pending_alignment = alignment
            i += 1
            continue

        table_meta = decode_table_meta_marker(line)
        if table_meta is not None:
            pending_table_meta = table_meta
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.strip() == EMPTY_PARAGRAPH_MARKER:
            blocks.append(
                {
                    "kind": "empty",
                    "style_meta": pending_paragraph_style,
                    "first_line_indent_cm": pending_first_line_indent_cm,
                    "alignment": pending_alignment,
                }
            )
            pending_paragraph_style = None
            pending_first_line_indent_cm = None
            pending_alignment = None
            pending_table_meta = None
            i += 1
            continue

        heading_match = heading_re.match(line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            blocks.append(
                {
                    "kind": "heading",
                    "level": level,
                    "text": text,
                    "style_meta": pending_paragraph_style,
                    "first_line_indent_cm": pending_first_line_indent_cm,
                    "alignment": pending_alignment,
                }
            )
            pending_paragraph_style = None
            pending_first_line_indent_cm = None
            pending_alignment = None
            pending_table_meta = None
            i += 1
            continue

        if line.strip().startswith("$$"):
            stripped = line.strip()
            if stripped == "$$":
                i += 1
                math_lines: List[str] = []
                while i < len(lines) and lines[i].strip() != "$$":
                    math_lines.append(lines[i])
                    i += 1
                if i < len(lines) and lines[i].strip() == "$$":
                    i += 1
                latex = "\n".join(math_lines).strip()
            else:
                single_line = re.match(r"^\s*\$\$(.*?)\$\$\s*$", line)
                if single_line:
                    latex = single_line.group(1).strip()
                else:
                    latex = stripped.strip("$").strip()
                i += 1

            payload = None
            if i < len(lines):
                payload_match = BLOCK_MARKER_RE.match(lines[i])
                if payload_match:
                    payload = payload_match.group(1)
                    i += 1

            blocks.append(
                {
                    "kind": "math",
                    "latex": latex,
                    "omml_payload": payload,
                    "style_meta": pending_paragraph_style,
                    "first_line_indent_cm": pending_first_line_indent_cm,
                    "alignment": pending_alignment,
                }
            )
            pending_paragraph_style = None
            pending_first_line_indent_cm = None
            pending_alignment = None
            pending_table_meta = None
            continue

        if is_markdown_table_start(lines, i):
            table_rows: List[List[str]] = [split_markdown_table_cells(lines[i])]
            i += 2
            while i < len(lines):
                next_line = lines[i]
                if not next_line.strip():
                    break
                if next_line.strip() == EMPTY_PARAGRAPH_MARKER:
                    break
                if heading_re.match(next_line):
                    break
                if next_line.strip().startswith("$$"):
                    break
                if not looks_like_markdown_table_row(next_line):
                    break
                if TABLE_SEPARATOR_RE.match(next_line):
                    break
                table_rows.append(split_markdown_table_cells(next_line))
                i += 1

            table_rows = normalize_markdown_table_rows(table_rows)
            blocks.append({"kind": "table", "rows": table_rows, "table_meta": pending_table_meta})
            pending_paragraph_style = None
            pending_first_line_indent_cm = None
            pending_alignment = None
            pending_table_meta = None
            continue

        paragraph_lines = [line]
        i += 1
        while i < len(lines):
            next_line = lines[i]
            if not next_line.strip():
                break
            if next_line.strip() == EMPTY_PARAGRAPH_MARKER:
                break
            if heading_re.match(next_line):
                break
            if next_line.strip().startswith("$$"):
                break
            if is_markdown_table_start(lines, i):
                break
            paragraph_lines.append(next_line)
            i += 1

        paragraph_text = "\n".join(paragraph_lines).strip()
        if paragraph_text:
            blocks.append(
                {
                    "kind": "paragraph",
                    "text": paragraph_text,
                    "style_meta": pending_paragraph_style,
                    "first_line_indent_cm": pending_first_line_indent_cm,
                    "alignment": pending_alignment,
                }
            )
            pending_paragraph_style = None
            pending_first_line_indent_cm = None
            pending_alignment = None
            pending_table_meta = None

    return blocks


def append_omml_to_paragraph(paragraph, payload: str) -> bool:
    xml_data = decode_omml_payload(payload)
    if not xml_data:
        return False

    try:
        element = parse_xml(xml_data)
    except Exception:
        try:
            parsed = etree.fromstring(xml_data)
            element = parse_xml(etree.tostring(parsed, encoding="utf-8"))
        except Exception:
            return False

    paragraph._p.append(element)
    return True


def clear_document_body(doc: DocxDocument) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != wtag("sectPr"):
            body.remove(child)


def find_heading_style(doc: DocxDocument, level: int) -> Optional[ParagraphStyle]:
    level_str = str(level)
    for style in doc.styles:
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue
        style_name = (style.name or "").lower()
        style_id = (style.style_id or "").lower()
        if re.search(rf"(heading|标题)\s*{level_str}\b", style_name):
            return cast(ParagraphStyle, style)
        if re.search(rf"(heading|标题)\s*{level_str}\b", style_id):
            return cast(ParagraphStyle, style)

    for candidate in (f"Heading {level}", f"heading {level}", f"标题 {level}"):
        try:
            return cast(ParagraphStyle, doc.styles[candidate])
        except Exception:
            pass

    return None


def copy_template_style_assets(template_docx: Path, output_docx: Path) -> None:
    style_assets = {
        "word/styles.xml",
        "word/fontTable.xml",
        "word/theme/theme1.xml",
    }

    with ZipFile(template_docx, "r") as template_zip:
        replacements = {
            name: template_zip.read(name)
            for name in style_assets
            if name in template_zip.namelist()
        }

    if not replacements:
        return

    temp_output = output_docx.with_suffix(output_docx.suffix + ".tmp")
    with ZipFile(output_docx, "r") as src_zip, ZipFile(temp_output, "w") as dst_zip:
        for info in src_zip.infolist():
            data = replacements.get(info.filename)
            if data is None:
                data = src_zip.read(info.filename)
            dst_zip.writestr(info, data)

    temp_output.replace(output_docx)


def markdown_to_docx(
    input_md: Path,
    output_docx: Path,
    prefer_omml: bool = True,
    template_docx: Optional[Path] = None,
) -> None:
    markdown_text = input_md.read_text(encoding="utf-8")
    blocks = parse_markdown_blocks(markdown_text)
    default_run_font = infer_default_run_font(template_docx) if template_docx is not None else None

    if template_docx is not None:
        doc = Document(str(template_docx))
        clear_document_body(doc)
    else:
        doc = Document()

    for block in blocks:
        kind = block["kind"]
        paragraph_style = find_paragraph_style_by_meta(doc, block.get("style_meta"))
        first_line_indent_cm = block.get("first_line_indent_cm")
        alignment = block.get("alignment")

        if kind == "heading":
            paragraph = doc.add_paragraph(block["text"])
            level = block["level"]
            if 1 <= level <= 6:
                style = find_heading_style(doc, level)
                if style is not None:
                    paragraph.style = style
                else:
                    paragraph.style = f"Heading {level}"
            else:
                paragraph.style = "Normal"
            if isinstance(alignment, str) and alignment:
                apply_paragraph_alignment(paragraph, alignment)
            continue

        if kind == "math":
            paragraph = doc.add_paragraph()
            if paragraph_style is not None:
                paragraph.style = paragraph_style
            if isinstance(first_line_indent_cm, (int, float)):
                apply_first_line_indent(paragraph, float(first_line_indent_cm))
            if isinstance(alignment, str) and alignment:
                apply_paragraph_alignment(paragraph, alignment)
            payload = block.get("omml_payload")
            restored = False
            if prefer_omml and payload:
                restored = append_omml_to_paragraph(paragraph, payload)
            if not restored:
                restored = append_latex_math_to_paragraph(paragraph, block["latex"], block=True)
            if not restored:
                run = paragraph.add_run(f"$$ {block['latex']} $$")
                apply_font_spec_to_run(run, default_run_font)
            continue

        if kind == "empty":
            paragraph = doc.add_paragraph()
            if paragraph_style is not None:
                paragraph.style = paragraph_style
            if isinstance(first_line_indent_cm, (int, float)):
                apply_first_line_indent(paragraph, float(first_line_indent_cm))
            if isinstance(alignment, str) and alignment:
                apply_paragraph_alignment(paragraph, alignment)
            continue

        if kind == "table":
            rows = block.get("rows", [])
            if not rows:
                continue

            table_meta = block.get("table_meta")
            row_count = len(rows)
            col_count = max(len(row) for row in rows)
            if isinstance(table_meta, dict):
                meta_row_count = table_meta.get("row_count")
                meta_col_count = table_meta.get("col_count")
                if isinstance(meta_row_count, int):
                    row_count = max(row_count, meta_row_count)
                if isinstance(meta_col_count, int):
                    col_count = max(col_count, meta_col_count)

            table = doc.add_table(rows=row_count, cols=col_count)
            apply_table_meta(table, table_meta if isinstance(table_meta, dict) else None)
            if not (isinstance(table_meta, dict) and table_meta.get("style_id")):
                try:
                    table.style = "Table Grid"
                except Exception:
                    pass

            row_metas = table_meta.get("rows_meta") if isinstance(table_meta, dict) else None
            if isinstance(row_metas, list):
                for row_idx, row_meta in enumerate(row_metas[:row_count]):
                    if isinstance(row_meta, dict):
                        apply_row_meta(table.rows[row_idx], row_meta)

            merges = table_meta.get("merges") if isinstance(table_meta, dict) else None
            normalized_merges: List[Dict[str, int]] = []
            if isinstance(merges, list):
                for merge in merges:
                    if not isinstance(merge, dict):
                        continue
                    row_idx = merge.get("row")
                    col_idx = merge.get("col")
                    rowspan = merge.get("rowspan")
                    colspan = merge.get("colspan")
                    if all(isinstance(value, int) for value in (row_idx, col_idx, rowspan, colspan)):
                        normalized_merges.append(cast(Dict[str, int], merge))

            for merge in normalized_merges:
                anchor = table.cell(merge["row"], merge["col"])
                target = table.cell(merge["row"] + merge["rowspan"] - 1, merge["col"] + merge["colspan"] - 1)
                anchor.merge(target)

            covered_positions = build_covered_positions_from_merges(normalized_merges)
            cells_meta = table_meta.get("cells") if isinstance(table_meta, dict) else None

            for row_idx in range(row_count):
                row = rows[row_idx] if row_idx < len(rows) else []
                for col_idx in range(col_count):
                    if (row_idx, col_idx) in covered_positions:
                        continue
                    cell_text = row[col_idx] if col_idx < len(row) else ""
                    cell = table.cell(row_idx, col_idx)
                    cell_meta = cells_meta.get(cell_key(row_idx, col_idx)) if isinstance(cells_meta, dict) else None
                    if isinstance(cell_meta, dict):
                        apply_cell_meta(cell, cell_meta)
                    append_markdown_cell_content(cell, cell_text, prefer_omml, doc, default_run_font, cell_meta)
            continue

        if kind == "paragraph":
            paragraph = doc.add_paragraph()
            if paragraph_style is not None:
                paragraph.style = paragraph_style
            if isinstance(first_line_indent_cm, (int, float)):
                apply_first_line_indent(paragraph, float(first_line_indent_cm))
            if isinstance(alignment, str) and alignment:
                apply_paragraph_alignment(paragraph, alignment)
            append_markdown_inline_content(paragraph, block["text"], prefer_omml, doc, default_run_font)
            continue

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))

    if template_docx is not None:
        copy_template_style_assets(template_docx, output_docx)


def analyze_docx(docx_path: Path) -> dict:
    document_root, styles_root = read_xml_from_docx(docx_path)
    style_map = style_id_to_name(styles_root)

    with ZipFile(docx_path, "r") as archive:
        styles_bytes = archive.read("word/styles.xml")
        font_table_bytes = archive.read("word/fontTable.xml") if "word/fontTable.xml" in archive.namelist() else b""

    style_counter: Counter[str] = Counter()
    heading_counter: Counter[int] = Counter()
    run_font_counter: Counter[Tuple[Optional[str], Optional[str], Optional[str], Optional[str]]] = Counter()
    bold_text_parts: List[str] = []
    underline_text_parts: List[str] = []
    strike_text_parts: List[str] = []
    highlighted_run_parts: List[Tuple[str, Optional[str], Optional[str]]] = []
    character_style_counter: Counter[str] = Counter()
    table_meta_list = [
        extract_table_meta(table, style_map) or {}
        for table in xpath_elements(document_root, "/w:document/w:body/w:tbl")
    ]

    paragraphs = xpath_elements(document_root, "/w:document/w:body/w:p")
    for paragraph in paragraphs:
        style_ids = xpath_strings(paragraph, "./w:pPr/w:pStyle/@w:val")
        style_id = style_ids[0] if style_ids else ""
        style_name = style_map.get(style_id, style_id or "Normal")
        style_counter[style_name] += 1

        heading_level = heading_level_from_style(style_id, style_map)
        if heading_level is not None:
            heading_counter[heading_level] += 1

    runs = xpath_elements(document_root, "/w:document/w:body//w:r")
    for run in runs:
        text = run_to_text(run)
        if is_run_bold(run):
            if text:
                bold_text_parts.append(text)

        underline = run_underline_value(run)
        if underline and text:
            underline_text_parts.append(text)

        if run_strike_value(run) and text:
            strike_text_parts.append(text)

        background_color = run_background_color_value(run)
        highlight = run_highlight_value(run)
        if text and (background_color or highlight):
            highlighted_run_parts.append((text, background_color, highlight))

        character_style = run_character_style_value(run)
        if character_style:
            character_style_counter[character_style] += 1

        font_nodes = xpath_elements(run, "./w:rPr/w:rFonts")
        if not font_nodes:
            continue

        font_values = (
            font_nodes[0].get(wtag("ascii")),
            font_nodes[0].get(wtag("hAnsi")),
            font_nodes[0].get(wtag("eastAsia")),
            font_nodes[0].get(wtag("cs")),
        )
        if any(font_values):
            run_font_counter[font_values] += 1

    o_math = len(xpath_elements(document_root, "/w:document/w:body//m:oMath"))
    o_math_para = len(xpath_elements(document_root, "/w:document/w:body//m:oMathPara"))
    dominant_run_font = None
    if run_font_counter:
        values = run_font_counter.most_common(1)[0][0]
        dominant_run_font = {attr: value for attr, value in zip(RUN_FONT_ATTRS, values) if value}

    bold_text_sha256 = None
    if bold_text_parts:
        bold_text_sha256 = sha256("".join(bold_text_parts).encode("utf-8")).hexdigest()

    underline_text_sha256 = None
    if underline_text_parts:
        underline_text_sha256 = sha256("".join(underline_text_parts).encode("utf-8")).hexdigest()

    strike_text_sha256 = None
    if strike_text_parts:
        strike_text_sha256 = sha256("".join(strike_text_parts).encode("utf-8")).hexdigest()

    highlight_sha256 = None
    if highlighted_run_parts:
        highlight_sha256 = sha256(
            json.dumps(highlighted_run_parts, ensure_ascii=False, sort_keys=True).encode("utf-8")
        ).hexdigest()

    table_meta_sha256 = sha256(
        json.dumps(table_meta_list, ensure_ascii=False, sort_keys=True).encode("utf-8")
    ).hexdigest()

    return {
        "paragraph_count": len(paragraphs),
        "table_count": len(table_meta_list),
        "table_meta_sha256": table_meta_sha256,
        "style_counts": dict(style_counter),
        "heading_counts": {str(k): v for k, v in sorted(heading_counter.items())},
        "bold_run_count": len(bold_text_parts),
        "bold_text_sha256": bold_text_sha256,
        "underline_run_count": len(underline_text_parts),
        "underline_text_sha256": underline_text_sha256,
        "strike_run_count": len(strike_text_parts),
        "strike_text_sha256": strike_text_sha256,
        "highlight_run_count": len(highlighted_run_parts),
        "highlight_sha256": highlight_sha256,
        "character_style_counts": dict(character_style_counter),
        "dominant_run_font": dominant_run_font,
        "oMath_count": o_math,
        "oMathPara_count": o_math_para,
        "math_total": o_math + o_math_para,
        "styles_sha256": sha256(styles_bytes).hexdigest(),
        "fontTable_sha256": sha256(font_table_bytes).hexdigest() if font_table_bytes else None,
    }


def cmd_docx2md(args: argparse.Namespace) -> int:
    input_docx = Path(args.input_docx)
    output_md = Path(args.output_md)
    output_md.parent.mkdir(parents=True, exist_ok=True)
    docx_to_markdown(input_docx, output_md, embed_omml=not args.no_embed_omml)
    print(f"Converted DOCX to Markdown: {output_md}")
    return 0


def cmd_md2docx(args: argparse.Namespace) -> int:
    input_md = Path(args.input_md)
    output_docx = Path(args.output_docx)
    template_docx = Path(args.template_docx) if args.template_docx else None
    markdown_to_docx(
        input_md,
        output_docx,
        prefer_omml=not args.no_prefer_omml,
        template_docx=template_docx,
    )
    print(f"Converted Markdown to DOCX: {output_docx}")
    return 0


def cmd_roundtrip(args: argparse.Namespace) -> int:
    input_docx = Path(args.input_docx)
    prefix = Path(args.prefix) if args.prefix else input_docx.with_suffix("")

    md_path = Path(f"{prefix}.md")
    roundtrip_docx = Path(f"{prefix}.roundtrip.docx")
    roundtrip_md = Path(f"{prefix}.roundtrip.md")
    report_path = Path(f"{prefix}.roundtrip_report.json")

    md_path.parent.mkdir(parents=True, exist_ok=True)

    docx_to_markdown(input_docx, md_path, embed_omml=True)
    markdown_to_docx(md_path, roundtrip_docx, prefer_omml=True, template_docx=input_docx)
    docx_to_markdown(roundtrip_docx, roundtrip_md, embed_omml=False)

    src_stats = analyze_docx(input_docx)
    rt_stats = analyze_docx(roundtrip_docx)

    checks = {
        "heading_counts_equal": src_stats["heading_counts"] == rt_stats["heading_counts"],
        "style_counts_equal": src_stats["style_counts"] == rt_stats["style_counts"],
        "bold_text_equal": src_stats["bold_text_sha256"] == rt_stats["bold_text_sha256"],
        "underline_text_equal": src_stats["underline_text_sha256"] == rt_stats["underline_text_sha256"],
        "strike_text_equal": src_stats["strike_text_sha256"] == rt_stats["strike_text_sha256"],
        "highlight_runs_equal": src_stats["highlight_sha256"] == rt_stats["highlight_sha256"],
        "character_style_counts_equal": src_stats["character_style_counts"] == rt_stats["character_style_counts"],
        "dominant_run_font_equal": src_stats["dominant_run_font"] == rt_stats["dominant_run_font"],
        "table_meta_equal": src_stats["table_meta_sha256"] == rt_stats["table_meta_sha256"],
        "oMath_equal": src_stats["oMath_count"] == rt_stats["oMath_count"],
        "oMathPara_equal": src_stats["oMathPara_count"] == rt_stats["oMathPara_count"],
        "paragraph_count_close": abs(src_stats["paragraph_count"] - rt_stats["paragraph_count"]) <= 2,
        "styles_xml_identical": src_stats["styles_sha256"] == rt_stats["styles_sha256"],
        "font_table_identical": src_stats["fontTable_sha256"] == rt_stats["fontTable_sha256"],
    }

    report = {
        "input_docx": str(input_docx),
        "markdown_output": str(md_path),
        "roundtrip_docx": str(roundtrip_docx),
        "roundtrip_markdown": str(roundtrip_md),
        "source_stats": src_stats,
        "roundtrip_stats": rt_stats,
        "checks": checks,
        "success": all(checks.values()),
    }

    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"Roundtrip report: {report_path}")
    print(json.dumps(checks, ensure_ascii=False, indent=2))
    return 0 if report["success"] else 1


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Convert between DOCX and Markdown for fixed report styles.")
    subparsers = parser.add_subparsers(dest="command", required=True)

    p_docx2md = subparsers.add_parser("docx2md", help="Convert DOCX to Markdown")
    p_docx2md.add_argument("input_docx", help="Input DOCX path")
    p_docx2md.add_argument("output_md", help="Output Markdown path")
    p_docx2md.add_argument("--no-embed-omml", action="store_true", help="Do not embed OMML payload markers")
    p_docx2md.set_defaults(func=cmd_docx2md)

    p_md2docx = subparsers.add_parser("md2docx", help="Convert Markdown to DOCX")
    p_md2docx.add_argument("input_md", help="Input Markdown path")
    p_md2docx.add_argument("output_docx", help="Output DOCX path")
    p_md2docx.add_argument("--no-prefer-omml", action="store_true", help="Do not restore OMML even if markers exist")
    p_md2docx.add_argument("--template-docx", default=None, help="Template DOCX to preserve styles/fonts")
    p_md2docx.set_defaults(func=cmd_md2docx)

    p_roundtrip = subparsers.add_parser("roundtrip", help="Run DOCX -> MD -> DOCX roundtrip validation")
    p_roundtrip.add_argument("input_docx", help="Input DOCX path")
    p_roundtrip.add_argument(
        "--prefix",
        default=None,
        help="Output prefix. Defaults to input DOCX path without extension.",
    )
    p_roundtrip.set_defaults(func=cmd_roundtrip)

    return parser


def main(argv: Optional[Sequence[str]] = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)
    return args.func(args)


if __name__ == "__main__":
    sys.exit(main())
